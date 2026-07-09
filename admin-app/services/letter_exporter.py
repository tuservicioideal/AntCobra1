"""
Letter Exporter — Multi-format Export
========================================
Takes a rendered template (list of Segments) and exports to:
  - DOCX  (python-docx)
  - PDF   (fpdf2)
  - JPG   (fpdf2 → pymupdf/fitz render)

Public API:
    export_letter(template_text, client, numero_carta, output_dir,
                  formats, seccion, gestor_name, campaign_id, gestor_config)
        → dict {fmt: filepath, ...}

    export_all_letters(by_seccion, numero_carta, output_dir, formats,
                       gestores_info, campaign_id, gestor_config)
        → dict {files, total_letters, errors, output_dir}

    build_zip(file_paths, zip_path)  → zip_path
"""
from __future__ import annotations

import io
import os
import zipfile
from typing import Optional

from .template_engine import (
    Segment, parse_template, render_template, CARTA_NOMBRES,
)

# ── Red colours per carta ────────────────────────────────────────
_TITLE_COLORS = {
    1: (99,  102, 241),   # Indigo   E1-1
    2: (234,  88,  12),   # Orange   E1-2
    3: (220,  38,  38),   # Red-600  E2-1
    4: (185,  28,  28),   # Red-700  E2-2
    5: (127,  29,  29),   # Red-900  E3-1
}

# ═══════════════════════════════════════════════════════════════
# DOCX
# ═══════════════════════════════════════════════════════════════

def _export_docx(segments: list[Segment], numero_carta: int, output_path: str,
                 gestor_config: dict) -> str:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    RED = RGBColor(*_TITLE_COLORS.get(numero_carta, (220, 38, 38)))
    TEXT = RGBColor(30, 41, 59)
    GREY = RGBColor(100, 116, 139)
    RED_INLINE = RGBColor(220, 38, 38)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"           # type: ignore
    style.font.size = Pt(11)              # type: ignore
    style.font.color.rgb = TEXT           # type: ignore
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    def _add_runs(para, seg: Segment, default_size: float = 11,
                  default_color: Optional[RGBColor] = None):
        for run in seg.runs:
            r = para.add_run(run.text)
            r.font.size = Pt(default_size)
            r.font.color.rgb = default_color or TEXT
            if run.bold:
                r.bold = True
            if run.italic:
                r.italic = True
            if run.red:
                r.font.color.rgb = RED_INLINE

    for seg in segments:
        if seg.kind == "blank":
            doc.add_paragraph()
            continue

        if seg.kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(seg.plain)
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = RED
            r.underline = True
            continue

        if seg.kind == "body":
            p = doc.add_paragraph()
            _add_runs(p, seg)
            continue

        if seg.kind == "center":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, seg)
            continue

        if seg.kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, seg)
            continue

        if seg.kind == "firma":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(seg.plain)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = TEXT
            continue

        if seg.kind == "nota":
            p = doc.add_paragraph()
            _add_runs(p, seg, default_size=8, default_color=GREY)
            p.paragraph_format.space_before = Pt(8)
            continue

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
# PDF  (fpdf2)
# ═══════════════════════════════════════════════════════════════

def _export_pdf(segments: list[Segment], numero_carta: int, output_path: str,
                gestor_config: dict) -> str:
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("Instale fpdf2: pip install fpdf2")

    TR, TG, TB = _TITLE_COLORS.get(numero_carta, (220, 38, 38))

    class _PDF(FPDF):
        pass

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(25, 20, 25)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    W = pdf.w - pdf.l_margin - pdf.r_margin   # usable width

    def _write_runs(seg: Segment, size: float = 11,
                    default_r=30, default_g=41, default_b=59):
        """Write runs inline on the current line, then do a line break."""
        for run in seg.runs:
            style = ""
            if run.bold:
                style += "B"
            if run.italic:
                style += "I"
            pdf.set_font("Helvetica", style=style, size=size)
            if run.red:
                pdf.set_text_color(220, 38, 38)
            else:
                pdf.set_text_color(default_r, default_g, default_b)
            # multi_cell for wrapping
            pdf.multi_cell(0, 5.5, run.text, new_x="RIGHT", new_y="LAST")
        pdf.ln(1)
        pdf.set_text_color(default_r, default_g, default_b)
        pdf.set_font("Helvetica", size=size)

    def _write_runs_center(seg: Segment, size: float = 11):
        full_text = seg.plain
        style = "B" if any(r.bold for r in seg.runs) else ""
        pdf.set_font("Helvetica", style=style, size=size)
        # Check for red
        has_red = any(r.red for r in seg.runs)
        if has_red:
            pdf.set_text_color(220, 38, 38)
        else:
            pdf.set_text_color(30, 41, 59)
        pdf.multi_cell(W, 5.5, full_text, align="C", new_x="LMARGIN")
        pdf.set_text_color(30, 41, 59)
        pdf.set_font("Helvetica", size=11)

    for seg in segments:
        if seg.kind == "blank":
            pdf.ln(4)
            continue

        if seg.kind == "title":
            pdf.set_font("Helvetica", style="BU", size=16)
            pdf.set_text_color(TR, TG, TB)
            pdf.multi_cell(W, 8, seg.plain, align="C", new_x="LMARGIN")
            pdf.set_text_color(30, 41, 59)
            pdf.set_font("Helvetica", size=11)
            pdf.ln(2)
            continue

        if seg.kind == "body":
            pdf.set_x(pdf.l_margin)
            _write_runs(seg)
            continue

        if seg.kind == "center":
            _write_runs_center(seg)
            continue

        if seg.kind == "bullet":
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", size=11)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(W - 6, 5.5, "• " + seg.plain, align="L",
                           new_x="LMARGIN")
            continue

        if seg.kind == "firma":
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(W, 5.5, seg.plain, align="C", new_x="LMARGIN")
            pdf.set_font("Helvetica", size=11)
            continue

        if seg.kind == "nota":
            pdf.set_font("Helvetica", style="I", size=8)
            pdf.set_text_color(100, 116, 139)
            pdf.ln(2)
            pdf.multi_cell(W, 4, seg.plain, align="L", new_x="LMARGIN")
            pdf.set_text_color(30, 41, 59)
            pdf.set_font("Helvetica", size=11)
            continue

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    pdf.output(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
# JPG  (PDF → pymupdf → image)
# ═══════════════════════════════════════════════════════════════

def _export_jpg(segments: list[Segment], numero_carta: int, output_path: str,
                gestor_config: dict) -> str:
    # First generate PDF to a temp buffer
    import tempfile

    tmp_pdf = output_path.replace(".jpg", "_tmp.pdf")
    _export_pdf(segments, numero_carta, tmp_pdf, gestor_config)

    try:
        import fitz  # type: ignore  # pymupdf
        doc = fitz.open(tmp_pdf)
        page = doc[0]
        # 150 DPI rendering
        mat = fitz.Matrix(150 / 72, 150 / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        pix.save(output_path)
        doc.close()
    finally:
        try:
            os.remove(tmp_pdf)
        except OSError:
            pass

    return output_path


# ═══════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════

def export_letter(
    template_text: str,
    client: dict,
    numero_carta: int,
    output_dir: str,
    formats: list[str],          # e.g. ["docx", "pdf", "jpg"]
    seccion: str = "",
    gestor_name: str = "",
    campaign_id: str = "",
    gestor_config: dict | None = None,
    campaign_info: dict | None = None,
    filename_prefix: str = "",
) -> dict[str, str]:
    """
    Render a template for one client and export to all requested formats.

    Returns a dict {format: filepath}.
    """
    gc = gestor_config or {}

    # Render tags
    filled = render_template(template_text, client, gc, campaign_info)
    segments = parse_template(filled)

    # Build base filename
    nombre = (
        client.get("nombre_completo", "").strip()
        or client.get("nombres", "cliente")
    ).replace(" ", "_")[:30]
    carta_code = CARTA_NOMBRES.get(numero_carta, f"Carta{numero_carta}").split("—")[0].strip().replace(" ", "_")
    sec_label = seccion.rsplit("_", 1)[-1] if "_" in seccion else seccion
    codigo = str(client.get("codigo_cliente", "")).strip()
    code_tag = f"Cli{codigo}_" if codigo else ""
    base = f"{filename_prefix}{code_tag}{carta_code}_Sec{sec_label}_{nombre}"

    results: dict[str, str] = {}
    os.makedirs(output_dir, exist_ok=True)

    for fmt in formats:
        fmt = fmt.lower()
        path = os.path.join(output_dir, f"{base}.{fmt}")
        try:
            if fmt == "docx":
                _export_docx(segments, numero_carta, path, gc)
            elif fmt == "pdf":
                _export_pdf(segments, numero_carta, path, gc)
            elif fmt == "jpg":
                _export_jpg(segments, numero_carta, path, gc)
            results[fmt] = path
        except Exception as e:
            results[f"{fmt}_error"] = str(e)

    return results


def export_all_letters(
    by_seccion: dict,
    numero_carta: int,
    output_dir: str,
    formats: list[str],
    gestores_info: dict | None = None,
    campaign_id: str = "",
    gestor_config: dict | None = None,
    template_text: str = "",
    campaign_info: dict | None = None,
) -> dict:
    """
    Export letters for all sections.

    Returns summary dict: files, total_letters, errors, output_dir.
    """
    files: list[str] = []
    entries: list[dict] = []
    total = 0
    errors: list[str] = []

    for seccion, clients in sorted(by_seccion.items()):
        gestor_name = (gestores_info or {}).get(seccion, "")
        for client in clients:
            try:
                result = export_letter(
                    template_text=template_text,
                    client=client,
                    numero_carta=numero_carta,
                    output_dir=output_dir,
                    formats=formats,
                    seccion=seccion,
                    gestor_name=gestor_name,
                    campaign_id=campaign_id,
                    gestor_config=gestor_config,
                    campaign_info=campaign_info,
                )
                for fmt, path in result.items():
                    if not fmt.endswith("_error"):
                        files.append(path)
                        entries.append({
                            "path": path,
                            "format": fmt.lower(),
                            "seccion_key": seccion,
                            "cliente_id": client.get("cliente_id", client.get("id", "")),
                            "codigo_cliente": client.get("codigo_cliente", ""),
                            "nombre_completo": client.get("nombre_completo", ""),
                        })
                total += 1
            except Exception as e:
                errors.append(f"Sección {seccion} — {client.get('nombre_completo','?')}: {e}")

    return {
        "files": files,
        "total_letters": total,
        "total_files": len(files),
        "entries": entries,
        "errors": errors,
        "output_dir": output_dir,
    }


def export_seccion_combined_pdf(
    by_seccion: dict,
    numero_carta: int,
    output_path: str,
    gestores_info: dict | None = None,
    campaign_id: str = "",
    gestor_config: dict | None = None,
    template_text: str = "",
    campaign_info: dict | None = None,
) -> str:
    """
    Generate one combined PDF with all clients, one page per client.
    Uses fpdf2 — appends pages for every client.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("Instale fpdf2: pip install fpdf2")

    gc = gestor_config or {}
    TR, TG, TB = _TITLE_COLORS.get(numero_carta, (220, 38, 38))

    class _PDF(FPDF):
        pass

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(25, 20, 25)
    pdf.set_auto_page_break(auto=True, margin=20)
    W = pdf.w - 25 - 25

    first_page = True
    for seccion, clients in sorted(by_seccion.items()):
        for client in clients:
            if not first_page:
                pdf.add_page()
            else:
                pdf.add_page()
                first_page = False

            filled = render_template(template_text, client, gc, campaign_info)
            segs = parse_template(filled)

            for seg in segs:
                if seg.kind == "blank":
                    pdf.ln(4)
                elif seg.kind == "title":
                    pdf.set_font("Helvetica", style="BU", size=16)
                    pdf.set_text_color(TR, TG, TB)
                    pdf.multi_cell(W, 8, seg.plain, align="C", new_x="LMARGIN")
                    pdf.set_text_color(30, 41, 59)
                    pdf.set_font("Helvetica", size=11)
                    pdf.ln(2)
                elif seg.kind == "body":
                    pdf.set_font("Helvetica", size=11)
                    pdf.set_text_color(30, 41, 59)
                    # Inline bold/italic/red via character-level calls
                    for run in seg.runs:
                        style = ("B" if run.bold else "") + ("I" if run.italic else "")
                        pdf.set_font("Helvetica", style=style, size=11)
                        if run.red:
                            pdf.set_text_color(220, 38, 38)
                        else:
                            pdf.set_text_color(30, 41, 59)
                        pdf.multi_cell(0, 5.5, run.text, new_x="RIGHT", new_y="LAST")
                    pdf.ln(1)
                elif seg.kind == "center":
                    pdf.set_font("Helvetica", style="B" if any(r.bold for r in seg.runs) else "", size=11)
                    has_red = any(r.red for r in seg.runs)
                    pdf.set_text_color(220, 38, 38 if has_red else 30)
                    pdf.multi_cell(W, 5.5, seg.plain, align="C", new_x="LMARGIN")
                    pdf.set_text_color(30, 41, 59)
                elif seg.kind == "bullet":
                    pdf.set_font("Helvetica", size=11)
                    pdf.set_text_color(30, 41, 59)
                    pdf.multi_cell(W - 6, 5.5, "• " + seg.plain, align="L", new_x="LMARGIN")
                elif seg.kind == "firma":
                    pdf.set_font("Helvetica", style="B", size=11)
                    pdf.multi_cell(W, 5.5, seg.plain, align="C", new_x="LMARGIN")
                    pdf.set_font("Helvetica", size=11)
                elif seg.kind == "nota":
                    pdf.set_font("Helvetica", style="I", size=8)
                    pdf.set_text_color(100, 116, 139)
                    pdf.ln(2)
                    pdf.multi_cell(W, 4, seg.plain, align="L", new_x="LMARGIN")
                    pdf.set_text_color(30, 41, 59)
                    pdf.set_font("Helvetica", size=11)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    pdf.output(output_path)
    return output_path


def build_zip(file_paths: list[str], zip_path: str) -> str:
    """Create a ZIP archive containing all given file paths."""
    os.makedirs(os.path.dirname(zip_path) or ".", exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in file_paths:
            if os.path.isfile(fp):
                zf.write(fp, os.path.basename(fp))
    return zip_path


# ═══════════════════════════════════════════════════════════════
# Word-template based export (preserves watermarks / images)
# ═══════════════════════════════════════════════════════════════

def export_letter_from_word(
    template_path: str,
    client: dict,
    numero_carta: int,
    output_dir: str,
    seccion: str = "",
    gestor_name: str = "",
    gestor_phone: str = "",
    campaign_id: str = "",
    gestor_config: dict | None = None,
    campaign_info: dict | None = None,
    filename_prefix: str = "",
    formats: list[str] | None = None,
    dpi: int = 200,
) -> dict[str, str | list[str]]:
    """
    Fill *template_path* (.docx) with client data and export documents.

    Supported formats: ``docx``, ``pdf``, ``jpg``.
    Raises on conversion failure so the caller can collect the error.
    """
    from .word_template_engine import word_template_to_documents, word_template_to_jpg
    from .template_engine import CARTA_NOMBRES

    nombre = (
        client.get("nombre_completo", "").strip()
        or client.get("nombres", "cliente")
    ).replace(" ", "_")[:30]
    carta_code = (
        CARTA_NOMBRES.get(numero_carta, f"Carta{numero_carta}")
        .split("—")[0].strip()
        .replace(" ", "_")
    )
    sec_label = seccion.rsplit("_", 1)[-1] if "_" in seccion else seccion
    codigo = str(client.get("codigo_cliente", "")).strip()
    code_tag = f"Cli{codigo}_" if codigo else ""
    base = f"{filename_prefix}{code_tag}{carta_code}_Sec{sec_label}_{nombre}"

    os.makedirs(output_dir, exist_ok=True)
    requested = [str(fmt).lower() for fmt in (formats or ["docx", "pdf"])]
    requested = [fmt for fmt in requested if fmt in {"docx", "pdf", "jpg"}]
    if not requested:
        raise ValueError("Debe solicitar al menos un formato válido.")

    result: dict[str, str | list[str]] = {}
    result.update(
        word_template_to_documents(
            template_path=template_path,
            client=client,
            output_dir=output_dir,
            base_filename=base,
            gestor_config=gestor_config,
            campaign_info=campaign_info,
            formats=requested,
            dpi=dpi,
            gestor_name=gestor_name,
            gestor_phone=gestor_phone,
        )
    )
    return result


def export_all_letters_from_word(
    by_seccion: dict,
    numero_carta: int,
    template_path: str,
    output_dir: str,
    gestores_info: dict | None = None,
    gestores_phones: dict | None = None,
    campaign_id: str = "",
    gestor_config: dict | None = None,
    campaign_info: dict | None = None,
    formats: list[str] | None = None,
    dpi: int = 200,
) -> dict:
    """
    Export documents for all sections using a Word template.

    Returns summary dict: files, total_letters, errors, output_dir.
    """
    files: list[str] = []
    entries: list[dict] = []
    total = 0
    errors: list[str] = []

    for seccion, clients in sorted(by_seccion.items()):
        gestor_name = (gestores_info or {}).get(seccion, "")
        gestor_phone = (gestores_phones or {}).get(seccion, "")
        for client in clients:
            client_label = client.get("nombre_completo", "?")
            try:
                result = export_letter_from_word(
                    template_path=template_path,
                    client=client,
                    numero_carta=numero_carta,
                    output_dir=output_dir,
                    seccion=seccion,
                    gestor_name=gestor_name,
                    gestor_phone=gestor_phone,
                    campaign_id=campaign_id,
                    gestor_config=gestor_config,
                    campaign_info=campaign_info,
                    formats=formats,
                    dpi=dpi,
                )
                client_files = 0
                for fmt, value in result.items():
                    if fmt.endswith("_error"):
                        errors.append(
                            f"Sección {seccion} — {client_label} ({fmt.replace('_error', '').upper()}): {value}"
                        )
                        continue
                    if fmt == "jpg" and isinstance(value, list):
                        for path in value:
                            files.append(path)
                            entries.append({
                                "path": path,
                                "format": "jpg",
                                "seccion_key": seccion,
                                "cliente_id": client.get("cliente_id", client.get("id", "")),
                                "codigo_cliente": client.get("codigo_cliente", ""),
                                "nombre_completo": client.get("nombre_completo", ""),
                            })
                            client_files += 1
                    elif isinstance(value, str):
                        files.append(value)
                        entries.append({
                            "path": value,
                            "format": fmt.lower(),
                            "seccion_key": seccion,
                            "cliente_id": client.get("cliente_id", client.get("id", "")),
                            "codigo_cliente": client.get("codigo_cliente", ""),
                            "nombre_completo": client.get("nombre_completo", ""),
                        })
                        client_files += 1
                if client_files > 0:
                    total += 1
                elif not any(k.endswith("_error") for k in result):
                    errors.append(
                        f"Sección {seccion} — {client_label}: no se generó ningún archivo."
                    )
            except Exception as e:
                errors.append(
                    f"Sección {seccion} — {client_label}: {e}"
                )

    return {
        "files": files,
        "total_letters": total,
        "total_files": len(files),
        "entries": entries,
        "errors": errors,
        "output_dir": output_dir,
    }
