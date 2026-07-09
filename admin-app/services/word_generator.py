from __future__ import annotations

"""
Word Document Generator — Cartas de Cobranza (5 Etapas)
=========================================================
Generates Word (.docx) files for debt collection letters.

Each carta corresponds to a collection etapa:
  - Carta 1 — E1-1 (Día  1,  Etapa 1): Invitación a Reingreso
  - Carta 2 — E1-2 (Día  9,  Etapa 1): No Pierdas Ser Empresaria
  - Carta 3 — E2-1 (Día 11,  Etapa 2): Requerimiento de Pago
  - Carta 4 — E2-2 (Día 35,  Etapa 2): Insistimos en el Pago
  - Carta 5 — E3-1 (Día 44,  Etapa 3): Exigimos Pago — Pre Judicial

Also includes:
  - generate_final_report() for Day-60 campaign summary
"""

import os
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT


def _seccion_display(seccion_key: str) -> str:
    """Extract plain section letter from a composite key like '01_1211_H' → 'H'."""
    if "_" in seccion_key:
        return seccion_key.rsplit("_", 1)[-1]
    return seccion_key

# ── Carta metadata per number ───────────────────────────────────
# 5 letters matching the real collection stages E1-1 … E3-1
CARTA_CONFIG = {
    # E1-1: Día 1 — Invitación a Reingreso (tono amigable/motivacional)
    1: {
        "codigo": "E1-1",
        "etapa": "Etapa 1",
        "titulo": "INVITACIÓN A REINGRESO",
        "subtitulo": "E1-1 — Primer Contacto",
        "asunto": "INVITACIÓN A REGULARIZAR SU SITUACIÓN DE PAGO",
        "color_titulo": RGBColor(79, 70, 229),      # Indigo
        "color_asunto": RGBColor(37, 99, 235),      # Blue-600
        "saludo": (
            "Nos complace saludarte y a la vez comunicarte que, de acuerdo "
            "a nuestros registros, mantienes un saldo pendiente de pago. "
            "Te invitamos a ponerte al día con tu cuenta y así continuar "
            "disfrutando de todos los beneficios y oportunidades que "
            "nuestra empresa tiene para ti."
        ),
        "cierre": (
            "Sabemos que en ocasiones surgen imprevistos que dificultan el "
            "cumplimiento de las obligaciones. Por eso, te invitamos a "
            "acercarte a nuestro representante asignado para coordinar "
            "juntos la mejor alternativa de pago.\n\n"
            "Tu fidelidad es importante para nosotros. Aprovecha esta "
            "oportunidad y mantén activa tu relación comercial con nosotros."
        ),
    },
    # E1-2: Día 9 — No Pierdas Ser Empresaria (tono motivacional/urgente)
    2: {
        "codigo": "E1-2",
        "etapa": "Etapa 1",
        "titulo": "NO PIERDAS SER EMPRESARIA",
        "subtitulo": "E1-2 — Segundo Contacto",
        "asunto": "IMPORTANTE: MANTENER TU ESTATUS DE EMPRESARIA",
        "color_titulo": RGBColor(234, 88, 12),      # Orange-600
        "color_asunto": RGBColor(194, 65, 12),      # Orange-700
        "saludo": (
            "Nos dirigimos a usted nuevamente para recordarle que a la "
            "fecha aún se registra un saldo pendiente en su cuenta. "
            "Le informamos que de no regularizarse, podría verse afectada "
            "su condición como empresaria activa, perdiendo los beneficios "
            "y acceso a nuestros productos y servicios."
        ),
        "cierre": (
            "Aún está a tiempo de evitar la baja de su cuenta. Regularice "
            "su deuda y continúe generando ingresos con nosotros.\n\n"
            "Un representante de cobranza se comunicará con usted en los "
            "próximos días. Le pedimos brindar las facilidades para "
            "coordinar su pago y así mantener su estatus activo."
        ),
    },
    # E2-1: Día 11 — Requerimiento de Pago (tono formal)
    3: {
        "codigo": "E2-1",
        "etapa": "Etapa 2",
        "titulo": "REQUERIMIENTO DE PAGO",
        "subtitulo": "E2-1 — Requerimiento Formal",
        "asunto": "REQUERIMIENTO FORMAL DE PAGO DE OBLIGACIÓN VENCIDA",
        "color_titulo": RGBColor(220, 38, 38),      # Red-600
        "color_asunto": RGBColor(185, 28, 28),      # Red-700
        "saludo": (
            "Por medio de la presente y de conformidad con las gestiones de "
            "cobranza iniciadas, le REQUERIMOS de manera formal que proceda "
            "al pago inmediato del saldo pendiente que mantiene con nuestra "
            "empresa. No obstante las comunicaciones previas realizadas, a "
            "la fecha no hemos registrado ningún pago ni coordinación "
            "formal de su parte."
        ),
        "cierre": (
            "Le instamos a regularizar su deuda en el plazo más breve "
            "posible para evitar el inicio de acciones adicionales de "
            "cobranza. Nuestro gestor asignado lo visitará para coordinar "
            "el cobro correspondiente.\n\n"
            "De no obtener respuesta favorable, su caso será escalado a "
            "la siguiente etapa de gestión, con las consecuencias que "
            "ello implica."
        ),
    },
    # E2-2: Día 35 — Insistimos en el Pago (tono más firme)
    4: {
        "codigo": "E2-2",
        "etapa": "Etapa 2",
        "titulo": "INSISTIMOS EN EL PAGO",
        "subtitulo": "E2-2 — Segunda Advertencia",
        "asunto": "SEGUNDO REQUERIMIENTO — PAGO URGENTE E INMEDIATO",
        "color_titulo": RGBColor(185, 28, 28),      # Red-700
        "color_asunto": RGBColor(153, 27, 27),      # Red-800
        "saludo": (
            "Lamentamos comunicarle que, a pesar de las notificaciones "
            "y gestiones de cobranza realizadas, su obligación de pago "
            "continúa IMPAGA. Su expediente se encuentra en proceso de "
            "evaluación para el siguiente nivel de acción, por lo que "
            "le EXHORTAMOS a atender esta situación de forma INMEDIATA."
        ),
        "cierre": (
            "Le reiteramos que de persistir el incumplimiento, nos veremos "
            "en la obligación de proceder con las medidas de cobranza "
            "reforzada que correspondan, lo que podría incluir:\n"
            "  • Reporte a centrales de riesgo crediticio\n"
            "  • Inicio de proceso de cobranza prejudicial\n"
            "  • Cargos adicionales por gastos de gestión\n\n"
            "Esta es una oportunidad para resolver su situación de forma "
            "directa y evitar mayores consecuencias. Comuníquese con "
            "nuestro representante a la brevedad."
        ),
    },
    # E3-1: Día 44 — Exigimos Pago — Pre Judicial (máxima urgencia)
    5: {
        "codigo": "E3-1",
        "etapa": "Etapa 3",
        "titulo": "EXIGIMOS PAGO — AVISO PRE JUDICIAL",
        "subtitulo": "E3-1 — Última Instancia",
        "asunto": "AVISO PRE JUDICIAL — EXIGENCIA INMEDIATA DE PAGO",
        "color_titulo": RGBColor(127, 29, 29),      # Red-900
        "color_asunto": RGBColor(127, 29, 29),      # Red-900
        "saludo": (
            "Por medio de la presente, y habiendo agotado las instancias "
            "previas de cobranza extrajudicial sin obtener respuesta "
            "satisfactoria de su parte, le comunicamos que su caso ha sido "
            "ESCALADO a la etapa final de gestión.\n\n"
            "Le EXIGIMOS el pago inmediato e integral de la deuda pendiente. "
            "Esta es la ÚLTIMA comunicación antes del cierre de la "
            "gestión extrajudicial y el inicio de las acciones legales "
            "que sean pertinentes."
        ),
        "cierre": (
            "IMPORTANTE: Si no regulariza su deuda dentro del plazo "
            "indicado, su expediente será derivado a la instancia judicial, "
            "lo que podría acarrear:\n"
            "  • Demanda judicial por cobro de deuda\n"
            "  • Embargo de bienes\n"
            "  • Reporte negativo en centrales de riesgo (Infocorp)\n"
            "  • Incremento del monto adeudado por honorarios y costas "
            "procesales\n\n"
            "Evite estas consecuencias tomando acción INMEDIATA. "
            "Comuníquese HOY con el gestor de cobranza asignado."
        ),
    },
}


def _format_currency(value):
    """Format a numeric value as Peruvian Soles."""
    try:
        v = float(value or 0)
        return f"S/ {v:,.2f}"
    except (ValueError, TypeError):
        return "S/ 0.00"


def _format_date_es():
    """Return today's date in Spanish."""
    return datetime.now().strftime("%d de %B de %Y").replace(
        "January", "enero").replace("February", "febrero").replace(
        "March", "marzo").replace("April", "abril").replace(
        "May", "mayo").replace("June", "junio").replace(
        "July", "julio").replace("August", "agosto").replace(
        "September", "septiembre").replace("October", "octubre").replace(
        "November", "noviembre").replace("December", "diciembre")


def _setup_doc():
    """Create a Document with default styling."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'  # type: ignore[union-attr]
    style.font.size = Pt(11)  # type: ignore[union-attr]
    style.font.color.rgb = RGBColor(30, 41, 59)  # type: ignore[union-attr]
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def _add_header(doc, numero_carta: int = 1, gestor_config: dict | None = None):
    """Add a carta-specific header with optional company letterhead."""
    cfg = CARTA_CONFIG.get(numero_carta, CARTA_CONFIG[1])
    gc = gestor_config or {}

    nombre_empresa = gc.get("nombre_empresa", "").strip()
    ruc_empresa    = gc.get("ruc_empresa", "").strip()
    dir_empresa    = gc.get("direccion_empresa", "").strip()

    # ── Company letterhead (if configured) ──
    if nombre_empresa:
        p_emp = doc.add_paragraph()
        p_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_emp = p_emp.add_run(nombre_empresa.upper())
        run_emp.bold = True
        run_emp.font.size = Pt(14)
        run_emp.font.color.rgb = RGBColor(30, 41, 59)

        if ruc_empresa:
            p_ruc = doc.add_paragraph()
            p_ruc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ruc = p_ruc.add_run(f"RUC: {ruc_empresa}")
            run_ruc.font.size = Pt(9)
            run_ruc.font.color.rgb = RGBColor(100, 116, 139)

        if dir_empresa:
            p_dir = doc.add_paragraph()
            p_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_dir = p_dir.add_run(dir_empresa)
            run_dir.font.size = Pt(9)
            run_dir.font.color.rgb = RGBColor(100, 116, 139)

        # Separator line
        p_sep = doc.add_paragraph()
        p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sep = p_sep.add_run("─" * 60)
        run_sep.font.size = Pt(9)
        run_sep.font.color.rgb = RGBColor(203, 213, 225)

    # ── Letter title and etapa label ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cfg["titulo"])
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = cfg["color_titulo"]

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{cfg['codigo']} — {cfg['etapa']}  ·  {cfg['subtitulo']}")
    run2.bold = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 116, 139)

    if not nombre_empresa:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("Sistema de Gestión de Cobranzas — Reacudo Legal")
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(148, 163, 184)


def _add_client_letter(doc, client: dict, seccion: str,
                       numero_carta: int = 1,
                       gestor_name: str = "",
                       campaign_id: str = "",
                       today: str = "",
                       gestor_config: dict | None = None):
    """Add one complete letter page to the document."""
    cfg = CARTA_CONFIG.get(numero_carta, CARTA_CONFIG[1])
    gc = gestor_config or {}

    _add_header(doc, numero_carta, gestor_config=gc)

    # Date and reference
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = p_date.add_run(f"Fecha: {today}")
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(100, 116, 139)

    if campaign_id:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_ref = p_ref.add_run(
            f"Ref: {campaign_id} / Sección {_seccion_display(seccion)} / {cfg['codigo']}")
        run_ref.font.size = Pt(9)
        run_ref.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # ── Addressee ──
    nombre = (client.get("nombre_completo", "").strip()
              or f"{client.get('nombres', '')} "
                 f"{client.get('apellido_paterno', '')} "
                 f"{client.get('apellido_materno', '')}".strip())
    dni = client.get("numero_documento", "")
    direccion = client.get("direccion", "")
    distrito = client.get("distrito", "")
    provincia = client.get("provincia", "")
    departamento = client.get("departamento", "")
    telefono = client.get("telefono_movil", "")

    p_to = doc.add_paragraph()
    run_sr = p_to.add_run("Señor(a): ")
    run_sr.bold = True
    run_sr.font.size = Pt(12)
    run_name = p_to.add_run(nombre.upper())
    run_name.bold = True
    run_name.font.size = Pt(12)

    if dni:
        p_dni = doc.add_paragraph()
        run_l = p_dni.add_run("DNI: ")
        run_l.bold = True
        p_dni.add_run(str(dni))

    addr_parts = [x for x in [direccion, distrito, provincia, departamento] if x]
    if addr_parts:
        p_addr = doc.add_paragraph()
        run_l = p_addr.add_run("Dirección: ")
        run_l.bold = True
        p_addr.add_run(", ".join(addr_parts))

    if telefono:
        p_tel = doc.add_paragraph()
        run_l = p_tel.add_run("Teléfono: ")
        run_l.bold = True
        p_tel.add_run(str(telefono))

    doc.add_paragraph()

    # ── Subject ──
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run(f"ASUNTO: {cfg['asunto']}")
    run_sub.bold = True
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = cfg["color_asunto"]

    doc.add_paragraph()

    # ── Body saludo ──
    deuda_asignada  = _format_currency(client.get("importe_deuda_asignada", 0))
    deuda_pendiente = _format_currency(client.get("importe_deuda_pendiente", 0))
    dias_atraso     = client.get("dias_atraso", 0)

    doc.add_paragraph(
        f"Estimado(a) {nombre},\n\n{cfg['saludo']}"
    )

    # ── Debt table ──
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    rows_data = [
        ("Deuda Asignada",    deuda_asignada),
        ("Deuda Pendiente",   deuda_pendiente),
        ("Días de Atraso",    str(dias_atraso)),
        ("Código de Cliente", str(client.get("codigo_cliente", ""))),
        ("Tipo de Carta",     f"{cfg['codigo']} — {cfg['titulo']}"),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # ── Closing paragraph ──
    doc.add_paragraph(cfg["cierre"])
    doc.add_paragraph()

    # ── Gestor / Signature block ──
    nombre_gestor = gc.get("nombre_gestor", gestor_name).strip()
    cargo_gestor  = gc.get("cargo_gestor", "Gestor de Cobranza").strip()
    tel_gestor    = gc.get("telefono_gestor", "").strip()
    correo_gestor = gc.get("correo_gestor", "").strip()

    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.add_run("_" * 40)

    p_sign2 = doc.add_paragraph()
    p_sign2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sign2.add_run(nombre_gestor if nombre_gestor else "Gestor de Cobranza")
    run_s.bold = True
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = RGBColor(30, 41, 59)

    p_cargo = doc.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cargo = p_cargo.add_run(cargo_gestor)
    run_cargo.font.size = Pt(10)
    run_cargo.font.color.rgb = RGBColor(100, 116, 139)

    if tel_gestor or correo_gestor:
        contact_parts = []
        if tel_gestor:
            contact_parts.append(f"Tel: {tel_gestor}")
        if correo_gestor:
            contact_parts.append(correo_gestor)
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact = p_contact.add_run(" | ".join(contact_parts))
        run_contact.font.size = Pt(9)
        run_contact.font.color.rgb = RGBColor(100, 116, 139)

    if gestor_name and gestor_name != nombre_gestor:
        p_g = doc.add_paragraph()
        p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_g = p_g.add_run(
            f"Gestor asignado: {gestor_name} — Sección {_seccion_display(seccion)}")
        run_g.font.size = Pt(9)
        run_g.font.color.rgb = RGBColor(148, 163, 184)


# ── Public API: per-gestor document ──────────────────────────────

def generate_letters_for_gestor(seccion: str, clients: list,
                                 output_dir: str,
                                 gestor_name: str = "",
                                 campaign_id: str = "",
                                 numero_carta: int = 1,
                                 gestor_config: dict | None = None) -> str:
    """
    Generate a Word document with one collection letter per client
    for a given gestor (seccion).

    Args:
        seccion:       Section letter (e.g. 'A', 'B')
        clients:       List of client dicts
        output_dir:    Directory to save the .docx file
        gestor_name:   Name of the assigned gestor (field gestor)
        campaign_id:   Campaign identifier
        numero_carta:  Carta number (1-5) — E1-1 to E3-1
        gestor_config: Dict with company/gestor static data
                       (nombre_empresa, ruc_empresa, nombre_gestor,
                        cargo_gestor, telefono_gestor, correo_gestor,
                        direccion_empresa)

    Returns:
        Full path to the generated .docx file
    """
    doc = _setup_doc()
    today = _format_date_es()

    for idx, client in enumerate(clients):
        if idx > 0:
            doc.add_page_break()
        _add_client_letter(
            doc, client, seccion,
            numero_carta=numero_carta,
            gestor_name=gestor_name,
            campaign_id=campaign_id,
            today=today,
            gestor_config=gestor_config,
        )

    # Save
    os.makedirs(output_dir, exist_ok=True)
    safe_name = ""
    if gestor_name:
        safe_name = "_" + "".join(
            c for c in gestor_name if c.isalnum() or c in (" ", "_")
        ).strip().replace(" ", "_")
    carta_cfg = CARTA_CONFIG.get(numero_carta, CARTA_CONFIG[1])
    codigo = carta_cfg.get("codigo", f"Carta{numero_carta}")
    filename = f"{codigo}_Seccion_{_seccion_display(seccion)}{safe_name}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generate_all_letters(by_seccion: dict, output_dir: str,
                          gestores_info: dict | None = None,
                          campaign_id: str = "",
                          numero_carta: int = 1,
                          gestor_config: dict | None = None) -> dict:
    """
    Generate Word documents for ALL gestores.

    Args:
        by_seccion:    Dict of {seccion_letter: [client_dicts, ...]}
        output_dir:    Base directory to save .docx files
        gestores_info: Optional dict {seccion: gestor_name}
        campaign_id:   Campaign identifier
        numero_carta:  Carta number (1-5) — E1-1 to E3-1
        gestor_config: Company/gestor static data dict

    Returns:
        dict: { files, total_letters, total_files, errors, output_dir }
    """
    files = []
    total = 0
    errors = []

    for seccion, clients in sorted(by_seccion.items()):
        try:
            gestor_name = (gestores_info or {}).get(seccion, "")
            path = generate_letters_for_gestor(
                seccion=seccion,
                clients=clients,
                output_dir=output_dir,
                gestor_name=gestor_name,
                campaign_id=campaign_id,
                numero_carta=numero_carta,
                gestor_config=gestor_config,
            )
            files.append(path)
            total += len(clients)
        except Exception as e:
            errors.append(f"Sección {seccion}: {str(e)}")

    return {
        "files": files,
        "total_letters": total,
        "total_files": len(files),
        "errors": errors,
        "output_dir": output_dir,
    }


# ── Tramo-aware: generate from pending-letters list ──────────────

def generate_tramo_letters(pending_list: list,
                           output_dir: str,
                           gestores_info: dict | None = None,
                           campaign_id: str = "",
                           all_clients_by_id: dict | None = None,
                           gestor_config: dict | None = None) -> dict:
    """
    Generate letters grouped by carta number and section, using the
    pending-letters list from CampaignManager.get_pending_letters().

    Each item in pending_list has:
      cliente_id, codigo_cliente, nombre_completo, numero_carta, tramo, saldo

    If all_clients_by_id is provided (id → full client dict), the full
    address/phone info is included in each letter.  Otherwise only the
    minimal info from pending_list is used.

    Returns:
        dict: { files, total_letters, by_carta: {1: n, 2: n, ...}, errors }
    """
    # Group: {(numero_carta, seccion): [client_dicts]}
    grouped = defaultdict(list)
    for item in pending_list:
        nc = item.get("numero_carta", 1)
        # Get full client data if available
        if all_clients_by_id and item.get("cliente_id") in all_clients_by_id:
            client = dict(all_clients_by_id[item["cliente_id"]])
        else:
            client = {
                "codigo_cliente": item.get("codigo_cliente", ""),
                "nombre_completo": item.get("nombre_completo", ""),
                "importe_deuda_pendiente": item.get("saldo", 0),
            }
        seccion = client.get("seccion", "?")
        grouped[(nc, seccion)].append(client)

    files = []
    total = 0
    by_carta = defaultdict(int)
    errors = []

    for (nc, seccion), clients in sorted(grouped.items()):
        try:
            gestor_name = (gestores_info or {}).get(seccion, "")
            path = generate_letters_for_gestor(
                seccion=seccion,
                clients=clients,
                output_dir=output_dir,
                gestor_name=gestor_name,
                campaign_id=campaign_id,
                numero_carta=nc,
                gestor_config=gestor_config,
            )
            files.append(path)
            total += len(clients)
            by_carta[nc] += len(clients)
        except Exception as e:
            errors.append(f"Carta {nc} Sec {seccion}: {str(e)}")

    return {
        "files": files,
        "total_letters": total,
        "total_files": len(files),
        "by_carta": dict(by_carta),
        "errors": errors,
        "output_dir": output_dir,
    }


# ── Day-60 Final Report ─────────────────────────────────────────

def generate_final_report(
    campaign_id: str,
    campaign_name: str,
    dia_campana: int,
    resumen: dict,
    secciones_stats: list,
    alertas: list,
    output_dir: str,
) -> str:
    """
    Generate a Word document with the Day-60 final campaign report.

    Args:
        campaign_id:     Campaign identifier
        campaign_name:   Human-readable campaign name
        dia_campana:     Current campaign day (should be >= 60)
        resumen:         Dict with aggregated stats (from get_campaign_status)
                         Keys: total, visitados, no_visitados, pagados,
                               morosos, no_ubica, suplantacion,
                               pago_no_registrado, deuda_visitada, etc.
        secciones_stats: List of dicts, one per section:
                         {seccion, gestor, total, visitados, pagados, morosos, ...}
        alertas:         List of alert dicts (from get_alerts)
        output_dir:      Where to save the report

    Returns:
        Full path to the generated .docx file
    """
    doc = _setup_doc()
    today = _format_date_es()

    # ── Cover / title ──
    doc.add_paragraph()
    doc.add_paragraph()
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_t.add_run("INFORME FINAL DE CAMPAÑA")
    run_t.bold = True
    run_t.font.size = Pt(24)
    run_t.font.color.rgb = RGBColor(79, 70, 229)

    p_n = doc.add_paragraph()
    p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_n = p_n.add_run(campaign_name or campaign_id)
    run_n.bold = True
    run_n.font.size = Pt(16)
    run_n.font.color.rgb = RGBColor(30, 41, 59)

    p_d = doc.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_d.add_run(f"Día de campaña: {dia_campana} — {today}")
    run_d.font.size = Pt(12)
    run_d.font.color.rgb = RGBColor(100, 116, 139)

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sys = p_sys.add_run("Sistema de Gestión de Cobranzas — Reacudo Legal")
    run_sys.font.size = Pt(10)
    run_sys.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_page_break()

    # ── 1. Resumen Ejecutivo ──
    h1 = doc.add_heading("1. Resumen Ejecutivo", level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    total = resumen.get("total", 0)
    visitados = resumen.get("deuda_visitada", resumen.get("visitados", 0))
    cobertura = (visitados / total * 100) if total else 0

    doc.add_paragraph(
        f"La campaña \"{campaign_name}\" finalizó en el día {dia_campana} "
        f"con un total de {total} clientes asignados. "
        f"Se logró una cobertura de visitas del {cobertura:.1f}%."
    )

    # Summary table
    summary_table = doc.add_table(rows=9, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    summary_rows = [
        ("Total de Clientes", str(total)),
        ("Clientes Visitados", str(visitados)),
        ("No Visitados", str(resumen.get("no_visitados", 0))),
        ("Pagados / Compromisos", str(resumen.get("pagados", 0))),
        ("Morosos Confirmados", str(resumen.get("morosos", 0))),
        ("No Ubicados", str(resumen.get("no_ubica", 0))),
        ("Suplantación", str(resumen.get("suplantacion", 0))),
        ("Pago No Registrado", str(resumen.get("pago_no_registrado", 0))),
        ("Cobertura de Visitas", f"{cobertura:.1f}%"),
    ]
    for i, (label, value) in enumerate(summary_rows):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        for paragraph in summary_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # ── 2. Desglose por Sección ──
    h2 = doc.add_heading("2. Resultados por Sección / Gestor", level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    if secciones_stats:
        cols = ["Sección", "Gestor", "Total", "Visitados",
                "Pagados", "Morosos", "No Ubica", "Suplant.", "Pago NR"]
        n_cols = len(cols)
        sec_table = doc.add_table(
            rows=1 + len(secciones_stats), cols=n_cols)
        sec_table.style = 'Light Shading Accent 1'
        sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, col_name in enumerate(cols):
            cell = sec_table.rows[0].cells[j]
            cell.text = col_name
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for i, sec in enumerate(secciones_stats, start=1):
            sec_table.rows[i].cells[0].text = str(sec.get("seccion", ""))
            sec_table.rows[i].cells[1].text = str(sec.get("gestor", ""))
            sec_table.rows[i].cells[2].text = str(sec.get("total", 0))
            sec_table.rows[i].cells[3].text = str(sec.get("visitados", 0))
            sec_table.rows[i].cells[4].text = str(sec.get("pagados", 0))
            sec_table.rows[i].cells[5].text = str(sec.get("morosos", 0))
            sec_table.rows[i].cells[6].text = str(sec.get("no_ubica", 0))
            sec_table.rows[i].cells[7].text = str(sec.get("suplantacion", 0))
            sec_table.rows[i].cells[8].text = str(
                sec.get("pago_no_registrado", 0))
            for j in range(n_cols):
                for paragraph in sec_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph("No hay datos desglosados por sección disponibles.")

    doc.add_paragraph()

    # ── 3. Alertas Registradas ──
    h3 = doc.add_heading("3. Alertas Registradas", level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    if alertas:
        doc.add_paragraph(
            f"Se registraron {len(alertas)} alertas durante la campaña.")

        alert_table = doc.add_table(rows=1 + len(alertas), cols=5)
        alert_table.style = 'Light Shading Accent 1'
        alert_headers = ["Tipo", "Cliente", "Gestor", "Nota", "Estado"]
        for j, h in enumerate(alert_headers):
            cell = alert_table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for i, a in enumerate(alertas, start=1):
            alert_table.rows[i].cells[0].text = str(
                a.get("tipo_alerta", ""))
            alert_table.rows[i].cells[1].text = str(
                a.get("codigo_cliente", ""))
            alert_table.rows[i].cells[2].text = str(
                a.get("gestor_nombre", a.get("gestor_id", "")))
            alert_table.rows[i].cells[3].text = str(
                a.get("nota_gestor", ""))[:80]
            alert_table.rows[i].cells[4].text = str(
                a.get("estado_alerta", ""))
            for j in range(5):
                for paragraph in alert_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
    else:
        doc.add_paragraph("No se registraron alertas durante la campaña.")

    doc.add_paragraph()

    # ── 4. Conclusión ──
    h4 = doc.add_heading("4. Conclusión", level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    pagados = resumen.get("pagados", 0)
    morosos = resumen.get("morosos", 0)
    no_ubica = resumen.get("no_ubica", 0)
    suplant = resumen.get("suplantacion", 0)

    doc.add_paragraph(
        f"De los {total} clientes asignados, {pagados} regularizaron su "
        f"deuda o establecieron compromiso de pago ({(pagados/total*100) if total else 0:.1f}%). "
        f"Se identificaron {morosos} morosos confirmados, {no_ubica} clientes "
        f"no ubicados en la dirección registrada, y {suplant} casos de "
        f"posible suplantación de identidad."
    )
    doc.add_paragraph(
        "Se recomienda trasladar los casos morosos pendientes a la siguiente "
        "etapa de cobranza y actualizar las direcciones de los clientes "
        "no ubicados para futuras campañas."
    )

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(
        f"Documento generado automáticamente — {today}")
    run_foot.font.size = Pt(8)
    run_foot.font.color.rgb = RGBColor(148, 163, 184)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    safe_name = "".join(
        c for c in (campaign_name or campaign_id)
        if c.isalnum() or c in (" ", "_", "-")
    ).strip().replace(" ", "_")
    filename = f"Informe_Final_{safe_name}_Dia{dia_campana}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath
